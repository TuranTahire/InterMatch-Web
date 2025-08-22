"""
PDF ve DOCX dosya okuma yardımcı fonksiyonları
"""

import PyPDF2
import docx
import io
import re
from typing import Optional

def extract_text_from_pdf(file_content: bytes) -> Optional[str]:
    """
    PDF dosyasından metin çıkarır
    
    Args:
        file_content: PDF dosyasının byte içeriği
        
    Returns:
        Çıkarılan metin veya None
    """
    try:
        print(f"🔍 PDF okuma başlatılıyor...")
        
        # PDF dosyasını oku
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        print(f"📄 PDF sayfa sayısı: {len(pdf_reader.pages)}")
        
        text = ""
        # Tüm sayfaları oku
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            print(f"📝 Sayfa {i+1}: {len(page_text)} karakter")
            
        final_text = text.strip()
        print(f"✅ PDF okuma tamamlandı: {len(final_text)} karakter")
        print(f"📋 İlk 200 karakter: {final_text[:200]}...")
        
        return final_text
        
    except Exception as e:
        print(f"❌ PDF okuma hatası: {e}")
        return None

def extract_text_from_docx(file_content: bytes) -> Optional[str]:
    """
    DOCX dosyasından metin çıkarır
    
    Args:
        file_content: DOCX dosyasının byte içeriği
        
    Returns:
        Çıkarılan metin veya None
    """
    try:
        print(f"🔍 DOCX okuma başlatılıyor...")
        
        # DOCX dosyasını oku
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        
        print(f"📄 DOCX paragraf sayısı: {len(doc.paragraphs)}")
        
        text = ""
        # Tüm paragrafları oku
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Boş paragrafları atla
                text += paragraph.text + "\n"
                print(f"📝 Paragraf {i+1}: {len(paragraph.text)} karakter")
            
        final_text = text.strip()
        print(f"✅ DOCX okuma tamamlandı: {len(final_text)} karakter")
        print(f"📋 İlk 200 karakter: {final_text[:200]}...")
        
        return final_text
        
    except Exception as e:
        print(f"❌ DOCX okuma hatası: {e}")
        return None

def extract_text_from_file(file_content: bytes, file_extension: str) -> Optional[str]:
    """
    Dosya uzantısına göre metin çıkarır
    
    Args:
        file_content: Dosya içeriği
        file_extension: Dosya uzantısı (.pdf, .docx)
        
    Returns:
        Çıkarılan metin veya None
    """
    print(f"🔍 Dosya okuma başlatılıyor: {file_extension}")
    print(f"📊 Dosya boyutu: {len(file_content)} byte")
    
    file_extension = file_extension.lower()
    
    if file_extension == '.pdf':
        result = extract_text_from_pdf(file_content)
    elif file_extension == '.docx':
        result = extract_text_from_docx(file_content)
    else:
        print(f"❌ Desteklenmeyen dosya formatı: {file_extension}")
        return None
    
    if result:
        print(f"✅ Dosya okuma başarılı: {len(result)} karakter")
    else:
        print(f"❌ Dosya okuma başarısız")
    
    return result

def clean_text(text: str) -> str:
    """
    Metni temizler ve formatlar
    
    Args:
        text: Ham metin
        
    Returns:
        Temizlenmiş metin
    """
    if not text:
        print("⚠️ Temizlenecek metin boş")
        return ""
    
    print(f"🧹 Metin temizleme başlatılıyor: {len(text)} karakter")
    
    # Fazla boşlukları temizle
    text = ' '.join(text.split())
    
    # Satır sonlarını düzenle
    text = text.replace('\n', ' ').replace('\r', ' ')
    
    final_text = text.strip()
    print(f"✅ Metin temizleme tamamlandı: {len(final_text)} karakter")
    print(f"📋 Temizlenmiş metin (ilk 300 karakter): {final_text[:300]}...")
    
    return final_text
